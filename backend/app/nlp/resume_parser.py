import logging
import io
from typing import Optional

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        return "\n".join(text_parts).strip()
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        return ""


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs).strip()
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return ""


def parse_resume(file_bytes: bytes, content_type: str) -> str:
    content_type = (content_type or "").lower()
    if "pdf" in content_type:
        return extract_text_from_pdf(file_bytes)
    elif "docx" in content_type or "word" in content_type or "msword" in content_type:
        return extract_text_from_docx(file_bytes)
    else:
        # Try PDF first, then DOCX, then plain text
        text = extract_text_from_pdf(file_bytes)
        if text:
            return text
        text = extract_text_from_docx(file_bytes)
        if text:
            return text
        try:
            return file_bytes.decode("utf-8", errors="ignore").strip()
        except Exception:
            return ""